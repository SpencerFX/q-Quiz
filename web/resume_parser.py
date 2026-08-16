import re

# Best-effort resume text extraction + section parsing for the Profile
# Edit page's "Upload / Auto-fill" button. Resume layouts vary wildly,
# so this is heuristic (section-header + regex based), not a real NLP
# parser - it's meant to save typing, not be authoritative. Callers
# should expect to review/edit what comes out of it.

SECTION_ALIASES = {
    "experience": [
        "work experience", "employment history", "professional experience",
        "work history", "professional background", "career history",
        "relevant experience", "employment", "experience"
    ],
    "education": [
        "education & training", "education and training", "academic background",
        "academic qualifications", "academic history", "educational background",
        "education"
    ],
    "skills": [
        "technical skills", "core competencies", "skills & tools", "skills and tools",
        "technical proficiencies", "areas of expertise", "key skills", "competencies",
        "skills"
    ],
    "links": [
        "links", "contact", "portfolio", "profiles", "social", "online presence"
    ],
    # Not extracted for anything, but still needs to be recognised as a
    # header - otherwise an unrecognised section (eg "Certifications")
    # never closes the section above it, and Education/Experience end up
    # swallowing everything all the way to the end of the document.
    "stop": [
        "certifications", "certification", "awards", "honors", "honours",
        "publications", "projects", "volunteer", "volunteering",
        "references", "languages", "interests", "hobbies", "summary",
        "objective", "profile", "additional information"
    ]
}

MONTH_RE = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?"

# Covers "2018 - 2022", "May 2023 - April 2025", "Aug. 2021 - May 2023",
# and open-ended ranges like "May 2026 - " (group 2/3 both absent means
# "still ongoing", handled by the callers below).
YEAR_RANGE_RE = re.compile(
    r"(?:" + MONTH_RE + r"\s+)?\b((?:19|20)\d{2})\b\s*[-–]\s*"
    r"(?:(?:" + MONTH_RE + r"\s+)?\b((?:19|20)\d{2})\b|(present|current))?",
    re.IGNORECASE
)

# Fallback for a single standalone date with no range, eg a graduation
# date: "College: BS - Economics - State University (December 2014)".
SINGLE_YEAR_RE = re.compile(r"\b((?:19|20)\d{2})\b")

DEGREE_RE = re.compile(
    r"\b(Bachelor|Master|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|MBA|Ph\.?D\.?|Associate|Diploma|BEng|MEng)\b",
    re.IGNORECASE
)

SCHOOL_KEYWORD_RE = re.compile(r"\b(University|College|Institute|School)\b", re.IGNORECASE)

# "[Company - Location; Month Year - Month Year]" - a whole job's company,
# location and dates bundled onto one bracketed line, with the title on
# the line above it instead of glued together.
BRACKET_RE = re.compile(r"\[(.*?)\]")

# A short label line with nothing but a category name, eg "Programming
# Languages:" or "SLDC / Infra Tools:" - real content, not a skill itself.
LABEL_LINE_RE = re.compile(r"^[A-Za-z0-9 /&-]{2,30}:\s*$")

URL_RE = re.compile(
    r"(https?://[^\s,;()]+|(?:www\.)?(?:linkedin\.com|github\.com|gitlab\.com|twitter\.com|x\.com|behance\.net|medium\.com)[^\s,;()]*)",
    re.IGNORECASE
)

DOMAIN_LABELS = {
    "linkedin.com": "LinkedIn",
    "github.com": "GitHub",
    "gitlab.com": "GitLab",
    "twitter.com": "Twitter",
    "x.com": "X",
    "behance.net": "Behance",
    "medium.com": "Medium"
}


def extract_text(path, ext):

    if ext == "pdf":
        return _extract_pdf_text(path)

    if ext == "docx":
        return _extract_docx_text(path)

    return ""


def _extract_pdf_text(path):

    import PyPDF2

    parts = []

    with open(path, "rb") as f:
        reader = PyPDF2.PdfFileReader(f)
        for i in range(reader.getNumPages()):
            parts.append(reader.getPage(i).extractText())

    return "\n".join(parts)


def _extract_docx_text(path):

    import docx

    doc = docx.Document(path)

    parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def _collapse_letter_spacing(line):

    # Handles headers rendered with tracked/spaced-out letters, eg
    # "P R O F E S S I O N A L    E X P E R I E N C E" - PDF text
    # extraction turns letter-tracking into literal single-space gaps
    # between characters, usually with a wider gap between words. Tokens
    # split on a single space are almost all one character long for a
    # line like this; a normal sentence won't be.
    tokens = line.split(" ")

    letters = [t for t in tokens if t]

    if len(letters) < 4 or sum(1 for t in letters if len(t) == 1) / len(letters) < 0.8:
        return line

    words, current = [], ""

    for t in tokens:
        if t == "":
            if current:
                words.append(current)
                current = ""
            continue
        if len(t) == 1:
            current += t
        else:
            if current:
                words.append(current)
                current = ""
            words.append(t)

    if current:
        words.append(current)

    return " ".join(words)


def _normalize_header_candidate(line):

    trimmed = re.sub(r"[\s\-=_.:]+$", "", line.strip())

    return _collapse_letter_spacing(trimmed).strip().lower()


def _find_sections(lines):

    header_positions = []

    for idx, raw in enumerate(lines):

        candidate = _normalize_header_candidate(raw)

        if not candidate or len(candidate) > 40:
            continue

        squashed = candidate.replace(" ", "")

        for key, aliases in SECTION_ALIASES.items():
            matched = any(
                candidate == a or candidate.startswith(a + " ") or candidate.startswith(a + "-")
                or squashed.startswith(a.replace(" ", ""))
                for a in aliases
            )
            if matched:
                header_positions.append((idx, key))
                break

    header_positions.sort()

    sections = {}

    for i, (idx, key) in enumerate(header_positions):
        end = header_positions[i + 1][0] if i + 1 < len(header_positions) else len(lines)
        sections.setdefault(key, [])
        sections[key].extend(lines[idx + 1:end])

    return sections


def _group_blocks(lines):

    blocks, current = [], []

    for l in lines:
        if not l.strip():
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(l.strip())

    if current:
        blocks.append(current)

    return blocks


def _split_by_date_anchor(blocks, lines):

    # PDF text extraction frequently drops the blank lines that separate
    # one job/degree from the next, so _group_blocks above collapses an
    # entire section into a single block. If there are visibly more date
    # ranges in the raw section than blocks we found, re-split on lines
    # that themselves contain a date range - each one marks a new entry
    # starting right after the previous anchor.
    non_empty = [l.strip() for l in lines if l.strip()]

    total_dates = len(YEAR_RANGE_RE.findall(" ".join(non_empty)))

    if total_dates <= len(blocks):
        return blocks

    anchor_idxs = [i for i, l in enumerate(non_empty) if YEAR_RANGE_RE.search(l)]

    if len(anchor_idxs) < 2:
        return blocks

    new_blocks, start = [], 0

    for i, idx in enumerate(anchor_idxs):
        end = anchor_idxs[i + 1] if i + 1 < len(anchor_idxs) else len(non_empty)
        chunk = non_empty[start:end]
        if chunk:
            new_blocks.append(chunk)
        start = end

    return new_blocks if new_blocks else blocks


def _split_by_bracket_anchor(lines):

    # "[Company - Location; Month Year - Month Year]" lines are a much
    # more specific, less ambiguous anchor than a bare date when they're
    # present - and unlike a date, a title line almost always sits
    # directly *above* its bracket rather than being merged with it, so
    # each entry's start is "the bracket's line minus one", not the
    # bracket line itself.
    non_empty = [l.strip() for l in lines if l.strip()]

    bracket_idxs = [i for i, l in enumerate(non_empty) if BRACKET_RE.search(l)]

    if len(bracket_idxs) < 2:
        return None

    starts = [max(0, idx - 1) for idx in bracket_idxs]

    for i in range(1, len(starts)):
        if starts[i] <= starts[i - 1]:
            starts[i] = starts[i - 1] + 1

    blocks = []

    for i, s in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(non_empty)
        chunk = non_empty[s:end]
        if chunk:
            blocks.append(chunk)

    return blocks


def _smart_split(blocks, lines):

    bracket_blocks = _split_by_bracket_anchor(lines)

    if bracket_blocks and len(bracket_blocks) > len(blocks):
        return bracket_blocks

    return _split_by_date_anchor(blocks, lines)


def _date_range(text):

    m = YEAR_RANGE_RE.search(text)

    if m:
        start = m.group(1)
        end = m.group(2) or m.group(3) or "Present"
        return start, end.title() if end.lower() in ("present", "current") else end, m

    single = SINGLE_YEAR_RE.search(text)

    if single:
        return "", single.group(1), single

    return "", "", None


def _parse_skills(lines):

    skills, seen = [], set()

    for raw in lines:

        line = raw.strip()

        if not line or LABEL_LINE_RE.match(line):
            continue

        has_delimiter = re.search(r"[,;•|/]|\t| {2,}", line)

        parts = re.split(r"[,;•|/]|\t| {2,}", line) if has_delimiter else [line]

        for p in parts:
            s = p.strip(" -•\t")
            if not s or len(s) > 40 or len(s.split()) > 6:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            skills.append(s)

    return skills[:40]


def _label_for_url(url):

    host = re.sub(r"^https?://", "", url, flags=re.IGNORECASE).split("/")[0].lower()

    if host.startswith("www."):
        host = host[4:]

    for domain, label in DOMAIN_LABELS.items():
        if host == domain or host.endswith("." + domain):
            return label

    return host


def _parse_links(all_text):

    links, seen = [], set()

    for m in URL_RE.finditer(all_text):
        url = m.group(0).rstrip(").,;")
        if not url.lower().startswith("http"):
            url = "https://" + url
        key = url.lower()
        if key in seen:
            continue
        seen.add(key)
        links.append({"label": _label_for_url(url), "url": url})

    return links[:10]


def _split_degree_school(line, date_match):

    working = re.sub(r"\([^)]*\)", "", line)
    working = re.sub(r"^(college|university|school)\s*:\s*", "", working, flags=re.IGNORECASE)

    if date_match and date_match.group(0) in working:
        working = working.replace(date_match.group(0), "")

    working = working.strip(" ,-–:")

    parts = [p.strip() for p in re.split(r"[-–,]", working) if p.strip()]

    # Backstop for a date that split() cut apart before it could be
    # stripped as one piece (eg a bare trailing year with no match) -
    # a lone 4-digit part is never a degree or school name.
    parts = [p for p in parts if not re.fullmatch(r"(?:19|20)\d{2}", p)]

    if not parts:
        return "", ""

    school_part = next((p for p in parts if SCHOOL_KEYWORD_RE.search(p)), None)

    if school_part:
        degree = ", ".join(p for p in parts if p != school_part)
        return degree, school_part

    degree_part = next((p for p in parts if DEGREE_RE.search(p)), parts[0])
    rest = next((p for p in parts if p != degree_part), "")

    return degree_part, rest


def _parse_education(lines):

    entries = []

    blocks = _smart_split(_group_blocks(lines), lines)

    for block in blocks:

        block_text = " ".join(block)

        start_date, end_date, date_match = _date_range(block_text)

        single_line = block[0] if len(block) == 1 else None

        if single_line:
            degree_line, school_line = _split_degree_school(single_line, date_match)
        else:
            degree_line = next((l for l in block if DEGREE_RE.search(l)), "")
            school_line = next((l for l in block if l and l != degree_line), block[0] if block else "")

        if not school_line and not degree_line:
            continue

        entries.append({
            "school": school_line[:80],
            "degree": degree_line[:80],
            "startDate": start_date[:20],
            "endDate": end_date[:20]
        })

    return entries[:10]


def _parse_experience(lines):

    entries = []

    blocks = _smart_split(_group_blocks(lines), lines)

    for block in blocks:

        if not block:
            continue

        block_text = " ".join(block)

        start_date, end_date, date_match = _date_range(block_text)

        header_line = block[0]

        # "[Company - Location; Month Year - Month Year]" on its own line
        # (usually right under the title) - check the first few lines
        # rather than just the header, since the title is often alone.
        bracket_idx, company, location = None, "", ""

        for i, l in enumerate(block[:3]):
            m = BRACKET_RE.search(l)
            if m:
                main = m.group(1).split(";")[0]
                bracket_parts = re.split(r"[-–]", main, maxsplit=1)
                company = bracket_parts[0].strip()
                location = bracket_parts[1].strip() if len(bracket_parts) > 1 else ""
                bracket_idx = i
                break

        if company:
            title = header_line.strip()
            description = " ".join(block[bracket_idx + 1:])[:400]
        else:
            header_for_split = header_line
            if date_match and date_match.group(0) in header_line:
                header_for_split = header_line.replace(date_match.group(0), "").strip(" ,-–")

            title, company = header_for_split, ""
            for sep in [" at ", ", ", " - ", " – ", " | "]:
                if sep in header_for_split:
                    title, company = header_for_split.split(sep, 1)
                    break
            company = company.strip() or "Unknown"
            description = " ".join(block[1:])[:400]

        entries.append({
            "company": company[:80],
            "title": title.strip()[:80],
            "startDate": start_date[:20],
            "endDate": end_date[:20],
            "location": location[:80],
            "description": description
        })

    return entries[:10]


def parse_resume(text):

    lines = text.splitlines()

    sections = _find_sections(lines)

    return {
        "skills": _parse_skills(sections.get("skills", [])),
        "education": _parse_education(sections.get("education", [])),
        "experience": _parse_experience(sections.get("experience", [])),
        "links": _parse_links(text)
    }
